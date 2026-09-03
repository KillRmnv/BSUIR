"""
Document loading strategy: extracts plain text from files of various formats.

Dispatch by file extension. Each extractor returns (text, error):
  text  — extracted plain text ("" on failure)
  error — None on success, human-readable message on failure

Supported: .txt .md .log .csv .pdf .docx .html .htm .rtf
"""
import io
import re
import html as html_module
from typing import Tuple, Optional


def _extract_txt(data: bytes, filename: str) -> Tuple[str, Optional[str]]:
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(enc), None
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore"), None


def _extract_pdf(data: bytes, filename: str) -> Tuple[str, Optional[str]]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "", "PDF support requires 'pypdf' (pip install pypdf)"
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip(), None
    except Exception as e:
        return "", f"Failed to parse PDF: {e}"


def _extract_docx(data: bytes, filename: str) -> Tuple[str, Optional[str]]:
    try:
        from docx import Document
    except ImportError:
        return "", "DOCX support requires 'python-docx' (pip install python-docx)"
    try:
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts).strip(), None
    except Exception as e:
        return "", f"Failed to parse DOCX: {e}"


def _extract_html(data: bytes, filename: str) -> Tuple[str, Optional[str]]:
    text = data.decode("utf-8", errors="ignore")
    text = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", text)
    text = html_module.unescape(text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n", text)
    return text.strip(), None


def _extract_rtf(data: bytes, filename: str) -> Tuple[str, Optional[str]]:
    text = data.decode("utf-8", errors="ignore")
    # Remove control words and groups, keep plain text
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", text)
    text = text.replace("\\'", "'")
    text = re.sub(r"[{}]", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip(), None


EXTRACTORS = {
    "txt": _extract_txt,
    "md": _extract_txt,
    "markdown": _extract_txt,
    "log": _extract_txt,
    "csv": _extract_txt,
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "html": _extract_html,
    "htm": _extract_html,
    "rtf": _extract_rtf,
}

SUPPORTED_FORMATS = ", ".join(
    "." + ext for ext in ("txt", "md", "pdf", "docx", "html", "rtf", "csv", "log")
)


def extract_text_from_bytes(data: bytes, filename: str) -> Tuple[str, Optional[str]]:
    """Extract plain text from a file's bytes. Returns (text, error)."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    extractor = EXTRACTORS.get(ext)
    if extractor is None:
        return "", f"Unsupported file format '.{ext}'. Supported: {SUPPORTED_FORMATS}"
    try:
        return extractor(data, filename)
    except Exception as e:
        return "", f"Failed to process file: {e}"